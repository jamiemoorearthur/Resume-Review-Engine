from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_edited_cv_docx(
    cv_text: str,
    suggested_bullets: list[dict],
    missing_keywords: list[str],
    section_recommendations: list[str],
) -> bytes:
    doc = Document()

    # Tighten default margins
    for section in doc.sections:
        section.top_margin = Pt(48)
        section.bottom_margin = Pt(48)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    # ── Original CV ──────────────────────────────────────────────────────────
    heading = doc.add_heading("Your CV", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for line in cv_text.split("\n"):
        stripped = line.strip()
        p = doc.add_paragraph(stripped if stripped else "")
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ── AI Suggested Improvements ────────────────────────────────────────────
    heading2 = doc.add_heading("AI Suggested Improvements", level=1)
    heading2.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    intro = doc.add_paragraph(
        "Apply the improvements below to increase your CV score. "
        "Suggested bullets replace the originals listed."
    )
    intro.paragraph_format.space_after = Pt(10)

    if suggested_bullets:
        doc.add_heading("Improved Bullet Points", level=2)
        for item in suggested_bullets:
            original = item.get("original", "")
            improved = item.get("improved", "")

            if original:
                orig_p = doc.add_paragraph(style="List Bullet")
                orig_run = orig_p.add_run(f"Original:  {original}")
                orig_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                orig_run.font.size = Pt(10)

            imp_p = doc.add_paragraph(style="List Bullet")
            imp_run = imp_p.add_run(f"Improved:  {improved}")
            imp_run.font.bold = True
            imp_run.font.color.rgb = RGBColor(0x10, 0x70, 0x5A)
            imp_run.font.size = Pt(10)

            doc.add_paragraph("")

    if missing_keywords:
        _add_horizontal_rule(doc)
        doc.add_heading("Keywords to Add", level=2)
        kw_p = doc.add_paragraph(", ".join(missing_keywords))
        kw_p.paragraph_format.space_after = Pt(8)

    if section_recommendations:
        _add_horizontal_rule(doc)
        doc.add_heading("Section Recommendations", level=2)
        for rec in section_recommendations:
            doc.add_paragraph(rec, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_edited_cv_pdf(
    cv_text: str,
    suggested_bullets: list[dict],
    missing_keywords: list[str],
    section_recommendations: list[str],
) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(20, 20, 20)
    pdf.add_page()

    # ── Original CV ──────────────────────────────────────────────────────────
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(26, 26, 46)
    pdf.cell(0, 10, "Your CV", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    pdf.set_font("Helvetica", size=10)
    pdf.set_text_color(40, 40, 40)
    for line in cv_text.split("\n"):
        stripped = line.strip()
        pdf.multi_cell(0, 5, stripped if stripped else "", new_x="LMARGIN", new_y="NEXT")

    # ── AI Suggested Improvements ────────────────────────────────────────────
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(26, 26, 46)
    pdf.cell(0, 10, "AI Suggested Improvements", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    pdf.set_font("Helvetica", "I", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.multi_cell(
        0, 6,
        "Apply the improvements below to increase your CV score. "
        "Suggested bullets replace the originals listed.",
        new_x="LMARGIN", new_y="NEXT",
    )
    pdf.ln(4)

    if suggested_bullets:
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 8, "Improved Bullet Points", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        for item in suggested_bullets:
            original = item.get("original", "")
            improved = item.get("improved", "")

            if original:
                pdf.set_font("Helvetica", size=10)
                pdf.set_text_color(150, 150, 150)
                pdf.multi_cell(0, 5, f"  Original:  {original}", new_x="LMARGIN", new_y="NEXT")

            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(16, 112, 90)
            pdf.multi_cell(0, 5, f"  Improved:  {improved}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

    if missing_keywords:
        pdf.ln(4)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 8, "Keywords to Add", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", size=10)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(0, 6, ", ".join(missing_keywords), new_x="LMARGIN", new_y="NEXT")

    if section_recommendations:
        pdf.ln(4)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 8, "Section Recommendations", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", size=10)
        pdf.set_text_color(40, 40, 40)
        for rec in section_recommendations:
            pdf.multi_cell(0, 6, f"  - {rec}", new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())
